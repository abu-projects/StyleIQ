from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "al.html"
OUT = ROOT / "docs"
OUT.mkdir(exist_ok=True)

INK = "1B1716"
MUTED = "6F675F"
FAINT = "9F9387"
GOLD = "C89B45"
GOLD_SOFT = "FFF9ED"
LINE = "E6DED6"
WHITE = "FFFFFF"

SECTION_NAMES_EN = {
    "S": "Welcome & Walkthrough",
    "A": "Authentication & First Use",
    "B": "Import & Closet Creation",
    "C": "Closet",
    "D": "Today & Daily Styling",
    "E": "Outfit Actions",
    "F": "Style Studio / Canvas",
    "G": "Saved Looks",
    "H": "Style Twin & Virtual Try-On",
    "I": "Planner & Events",
    "J": "Trips",
    "K": "Discover & Community",
    "L": "Profile & Settings",
    "M": "Muse Stylist",
}

SECTION_NAMES_AR = {
    "S": "الترحيب والجولة التعريفية",
    "A": "الدخول والاستخدام الأول",
    "B": "إضافة القطع وبناء الخزانة",
    "C": "الخزانة",
    "D": "اليوم والتنسيق اليومي",
    "E": "إجراءات الإطلالة",
    "F": "استوديو التنسيق",
    "G": "الإطلالات المحفوظة",
    "H": "Style Twin والتجربة الافتراضية",
    "I": "المخطط والمناسبات",
    "J": "الرحلات",
    "K": "الاستكشاف والمجتمع",
    "L": "الملف الشخصي والإعدادات",
    "M": "مساعد Muse",
}

SECTION_ROLE_EN = {
    "S": "Introduces the StyleIQ proposition and lets the user enter the app or begin account setup.",
    "A": "Collects first-use account and preference signals while keeping optional personalization skippable.",
    "B": "Turns photos, searches, or receipts into confirmed Closet items with editable metadata.",
    "C": "Lets the user browse, inspect, edit, and style the items they own or have saved.",
    "D": "Produces context-aware daily outfit recommendations from weather, plans, and Closet data.",
    "E": "Captures the user's response to a Look and routes it into saving, editing, sharing, or try-on.",
    "F": "Provides the editable outfit canvas where pieces, sources, context, and layout are managed.",
    "G": "Stores complete Looks and exposes their pieces, metadata, value, visibility, and activity.",
    "H": "Configures the optional private Style Twin and renders a Look on the user's visual reference.",
    "I": "Connects dates and events to Closet items, saved Looks, and Muse-assisted planning.",
    "J": "Builds a trip from destinations and activities through packing, daily Looks, and sharing.",
    "K": "Supports discovery of Looks, creators, products, and community interactions.",
    "L": "Manages the user's profile, learned preferences, privacy, support, and application settings.",
    "M": "Routes styling intents to recommendations, explanations, or evidence-based wardrobe insights.",
}

SECTION_ROLE_AR = {
    "S": "يعرّف المستخدم بقيمة StyleIQ ويتيح له دخول التطبيق أو بدء إنشاء الحساب.",
    "A": "يجمع بيانات الحساب وإشارات التفضيل الأولى مع إبقاء خطوات التخصيص الاختيارية قابلة للتخطي.",
    "B": "يحوّل الصور أو البحث أو الإيصالات إلى قطع مؤكدة داخل الخزانة ببيانات قابلة للتعديل.",
    "C": "يتيح تصفح القطع المملوكة أو المحفوظة وفحصها وتعديلها وتنسيقها.",
    "D": "ينتج توصيات يومية حسب الطقس والخطط والبيانات الفعلية للخزانة.",
    "E": "يسجل رد فعل المستخدم تجاه الإطلالة ويوجهه للحفظ أو التعديل أو المشاركة أو التجربة.",
    "F": "يوفر مساحة تحرير الإطلالة وإدارة القطع والمصادر والسياق وترتيب الطبقات.",
    "G": "يحفظ الإطلالات الكاملة ويعرض قطعها وبياناتها وقيمتها وخصوصيتها ونشاطها.",
    "H": "يجهز Style Twin الاختياري والخاص ويعرض الإطلالة على المرجع البصري للمستخدم.",
    "I": "يربط الأيام والمناسبات بقطع الخزانة والإطلالات المحفوظة وتخطيط Muse.",
    "J": "يبني رحلة من الوجهات والأنشطة مرورًا بالتعبئة والإطلالات اليومية والمشاركة.",
    "K": "يدعم اكتشاف الإطلالات وصنّاع المحتوى والمنتجات والتفاعل المجتمعي.",
    "L": "يدير الملف الشخصي والتفضيلات المتعلمة والخصوصية والدعم وإعدادات التطبيق.",
    "M": "يوجه نية المستخدم إلى توصية أو تفسير أو تحليل مبني على بيانات الخزانة.",
}

SECTION_SHORT_EN = {
    "S": "Brand entry and first-use orientation.",
    "A": "Account setup and initial preference signals.",
    "B": "Item intake and Closet creation.",
    "C": "Browse, edit, and style owned items.",
    "D": "Daily recommendations based on context.",
    "E": "Save, edit, share, try on, or reject a Look.",
    "F": "Editable outfit canvas and piece management.",
    "G": "Saved Look library and metadata.",
    "H": "Style Twin setup and private rendering.",
    "I": "Calendar and event-based Look planning.",
    "J": "Trip, packing, and daily-Look planning.",
    "K": "Discover, search, and community actions.",
    "L": "Profile, preferences, privacy, and support.",
    "M": "Muse prompts, explanations, and wardrobe insights.",
}

SECTION_SHORT_AR = {
    "S": "دخول العلامة والجولة الأولى.",
    "A": "إعداد الحساب وإشارات التفضيل الأولى.",
    "B": "إدخال القطع وبناء الخزانة.",
    "C": "تصفح القطع وتعديلها وتنسيقها.",
    "D": "توصيات يومية حسب السياق.",
    "E": "حفظ الإطلالة أو تعديلها أو مشاركتها أو رفضها.",
    "F": "تحرير الإطلالة وإدارة القطع.",
    "G": "مكتبة الإطلالات وبياناتها.",
    "H": "إعداد Style Twin وعرض الإطلالات.",
    "I": "التقويم وتخطيط المناسبات.",
    "J": "الرحلات والتعبئة والإطلالات اليومية.",
    "K": "الاستكشاف والبحث والتفاعل.",
    "L": "الملف والتفضيلات والخصوصية والدعم.",
    "M": "أسئلة Muse وتفسيراته وتحليلات الخزانة.",
}

AR_DATA = {
    "S": [
        ("شاشة افتتاح StyleIQ", "مقدمة بصرية وهوية تحريرية تمثل أول نقطة دخول إلى التطبيق."),
        ("التعرّف على Muse", "تعريف المستخدم بالمساعد الشخصي وعرض مختصر لما يستطيع StyleIQ تقديمه."),
        ("لمن نقوم بالتنسيق؟", "اختيار نوع الخزانة التي سيقوم StyleIQ بتخصيصها قبل إعداد الحساب."),
    ],
    "A": [
        ("إنشاء حساب أو تسجيل الدخول", "الدخول باستخدام Google أو Apple أو البريد الإلكتروني."),
        ("إنشاء الحساب: الاسم", "إدخال الاسم الأول واسم العائلة مع دعم التعبئة المسبقة من OAuth."),
        ("التسجيل بالبريد الإلكتروني", "تسجيل بريد الحساب المستخدم للدخول والتعرّف على إيصالات التسوق التي يختار المستخدم تمريرها."),
        ("تأكيد البريد: رمز OTP", "إدخال رمز من ستة أرقام مع إعادة الإرسال وحالات الخطأ والتأكيد."),
        ("إعداد الملف: المهنة", "اختيار مهنة مقترحة أو كتابة مهنة مخصصة قابلة للبحث."),
        ("تفضيل نوع الخزانة", "اختيار ملابس نسائية أو رجالية أو الاثنين معًا."),
        ("مقدمة Style Twin", "شرح التجربة الافتراضية مع إمكانية المتابعة أو التخطي."),
        ("إعداد صور Style Twin", "إضافة صورة وجه وصورة جسم منفصلتين مع توضيح الخصوصية."),
        ("كيف عرفت عنا؟", "تسجيل مصدر الوصول مثل TikTok أو Instagram أو صديق أو ChatGPT أو Google."),
        ("تفضيلات الستايل: البراندات", "البحث عن البراندات واختيار ثلاث علامات على الأقل كنقطة بداية للتفضيلات."),
        ("شاشة Today لأول استخدام", "تشرح وعد البدء بقطعة واحدة وتعرض إجراء الإضافة الأساسي."),
        ("مقدمة استيراد الخزانة", "تعرض الفرق قبل وبعد معالجة Prettify وتشرح فائدة تحسين صور القطع."),
        ("الخزانة الفارغة", "حالة فارغة مطمئنة تدعو المستخدم إلى إضافة أول قطعة إلى الخزانة."),
        ("الإطلالات المحفوظة الفارغة", "تشرح معنى حفظ الإطلالة وما الذي سيبقى محفوظًا معها."),
        ("الرحلات الفارغة", "تشرح تخطيط الوجهة وقائمة التعبئة والإطلالات اليومية للرحلة."),
    ],
    "B": [
        ("قائمة الإضافة العامة", "تجمع طرق الإضافة: الصور، البحث عن قطعة، أو تمرير الإيصالات."),
        ("اختيار الصور", "اختيار صورة واحدة أو عدة صور للقطع مع تهيئتها للتحسين."),
        ("دليل Prettify", "يوضح تحويل صورة عادية إلى صورة قطعة نظيفة مناسبة للاستوديو."),
        ("البحث للاستيراد", "البحث بكلمة مفتاحية أو لصق رابط منتج."),
        ("نتائج البحث", "عرض صورة القطعة واسمها ومصدرها والبراند قبل الاختيار."),
        ("تأكيد القطعة", "مراجعة البيانات القابلة للتعديل وتشغيل Prettify ثم إضافة قطعة أو عدة قطع."),
        ("جاري التحسين", "حالة معالجة تمنع الإضافة مؤقتًا مع إبقاء البيانات المرئية للمراجعة."),
        ("نتيجة المعالجة", "عرض الصورة الناتجة مع خيارات إعادة المحاولة أو الإضافة."),
        ("تمرير الإيصالات", "عرض تعليمات شخصية لاستيراد المشتريات من البريد الإلكتروني."),
        ("مسودات الاستيراد", "إظهار عدد المسودات غير المكتملة والحفاظ عليها للعودة لاحقًا."),
        ("نجاح الإضافة", "تأكيد نجاح الإضافة وتحديث الخزانة وتحويل تجربة Today من الحالة الفارغة."),
    ],
    "C": [
        ("شبكة الخزانة", "عرض الخزانة وWishlist مع الفرز والفلاتر والتصنيفات."),
        ("تفاصيل القطعة", "عرض الصورة الأساسية والبيانات القابلة للتعديل للقطعة."),
        ("أدوات الصورة", "قص الصورة وتشغيل Prettify أو إعادة المحاولة."),
        ("البيانات الأساسية", "تعديل البراند والاسم والتصنيف والألوان والمواسم والمقاس والوسوم والسعر."),
        ("بيانات التنسيق", "تحديد قواعد اللبس والأطوال والخامات ومدى صلاحية القطعة للتنسيق."),
        ("ذكاء القطعة", "عرض الإطلالات المحفوظة وعدد مرات الارتداء والتكلفة لكل ارتداء."),
        ("نسّق هذه القطعة", "إنشاء عائلات مناسبات مع تثبيت القطعة المحددة داخل الإطلالة."),
    ],
    "D": [
        ("قبل وجود بيانات للخزانة", "حالة Today قبل إضافة قطع، وتوجه المستخدم إلى الاستيراد أولًا."),
        ("بعد إضافة أول قطعة", "يعرض الطقس والسياق وتوصية Muse اليومية بعد توفر بيانات الخزانة."),
        ("جاري إعداد التوصية", "حالة تحميل تصف عمل Muse بصياغة بشرية وواضحة."),
        ("عائلات الإطلالات", "عرض اتجاهات مثل Business casual والحفلات والرسمي والمهني وشبه الرسمي."),
        ("عارض البدائل", "التنقل بين عدة نسخ من الإطلالة داخل المناسبة نفسها."),
        ("تصنيف مفقود", "يوضح أن الإطلالة تحتاج نوع قطعة غير موجود ومتوافق في الخزانة."),
    ],
    "E": [
        ("الحفظ بالقلب", "حفظ الإطلالة الحالية مع سياقها وقطعها."),
        ("التعديل بالقلم", "فتح الإطلالة داخل Style Studio للتخصيص."),
        ("غير مناسبة لي", "تسجيل ملاحظة سلبية لتحسين التوصيات التالية."),
        ("الإرسال والمشاركة", "فتح تدفق مشاركة الإطلالة مع الحفاظ على البيانات الخاصة."),
        ("إنشاء Style Twin", "فتح التجربة الافتراضية أو بوابة إعداد Style Twin عند الحاجة."),
        ("السحب بين البدائل", "الانتقال بين بدائل الإطلالة وتحديث الاختيار الحالي."),
    ],
    "F": [
        ("هيدر الاستوديو", "يوفر الرجوع وتعديل اسم الإطلالة وحفظ المسودة."),
        ("Flat lay", "تركيب قطع بخلفيات منزوعة داخل مشهد مسطح قابل للتعديل."),
        ("الدخول إلى Style Twin", "إضافة العرض على Style Twin كتحسين اختياري وليس شرطًا."),
        ("سياق الإطلالة", "إدارة الملاحظات والتقويم والتاريخ والموقع المرتبطين بالإطلالة."),
        ("طبقة التصنيف", "عرض بدائل التصنيف مع الإخفاء والإظهار والإضافة."),
        ("إضافة طبقة", "فتح محدد المصدر لإضافة قطعة جديدة إلى الإطلالة."),
        ("المحدد اليدوي", "البحث والفلترة داخل Closet أو Wishlist واختيار القطعة."),
        ("تحديد عدة طبقات", "اختيار عدة قطع وإدارتها دفعة واحدة."),
        ("مصدر القطعة", "تمييز القطع المملوكة عن اقتراحات الشراء داخل الإطلالة."),
        ("قطعة مفقودة", "عرض مكان واضح للقطعة الناقصة مع تفسير سبب الحاجة إليها."),
        ("استمرار المسودة", "حفظ الاسم والقطع والسياق والترتيب لاستكمال العمل لاحقًا."),
    ],
    "G": [
        ("شبكة الإطلالات", "فرز الإطلالات وفلترتها وتحديدها وعرض Lookbooks وسجل الارتداء."),
        ("تفاصيل الإطلالة", "عرض المعاينة والوصف والتاريخ والموقع والملاحظات."),
        ("قطع الإطلالة", "تمييز القطع المملوكة والمقترحة وإظهار روابط التسوق عند الحاجة."),
        ("تصنيف الإطلالة", "عرض الوسوم وقواعد اللبس والمواسم والستايلات المرتبطة."),
        ("قيمة الإطلالة", "حساب القيمة المعروفة والتعامل بوضوح مع الأسعار غير المتوفرة."),
        ("الظهور والخصوصية", "اختيار خاص أو للمتابعين أو عام على Discover."),
        ("مؤشرات التفاعل", "عرض الإعجابات والتعليقات وإعادة التنسيق والإرسال الخاص."),
        ("حذف الإطلالة", "تأكيد إجراء الحذف مع توضيح أثره وعدم إمكانية التراجع."),
    ],
    "H": [
        ("مقدمة Style Twin", "تشرح عرض الإطلالات المحفوظة على مرجع بصري خاص."),
        ("إرشادات صورة الوجه", "توضح الحاجة إلى وجه واحد ظاهر وإضاءة متوازنة وعدم وجود شخص آخر."),
        ("رفع صورة الوجه", "اختيار الصورة وقصها ومعالجتها قبل الحفظ."),
        ("مرجع الجسم", "إضافة صورة جسم منفصلة لتحسين النسب في العرض."),
        ("القياسات", "إدخال الطول والوزن بوحدات مترية أو إمبراطورية."),
        ("المظهر", "اختيار لون البشرة والشعر كبيانات قابلة للتعديل."),
        ("الوضعية", "اختيار الوقوف أو الميل أو الهاتف أو المشي أو الوضع العشوائي."),
        ("تفضيل إدخال القميص", "اختيار Tucked أو Untucked أو بدون تفضيل."),
        ("حد الاستخدام", "عرض عدد مرات التوليد المتبقية وموعد إعادة الضبط."),
        ("نتيجة Style Twin", "عرض الإطلالة الناتجة مع الحفظ وإعادة المحاولة وإرسال الملاحظات."),
    ],
    "I": [
        ("تقويم الشهر", "التنقل بين الشهور وإضافة خطة ليوم محدد وإبراز اليوم الحالي."),
        ("لوحة تقدم الستايل", "عرض الاستمرارية وأكثر القطع ارتداءً واستخدام الخزانة ومسار التطور."),
        ("قائمة اليوم", "اختيار قطعة من الخزانة أو إطلالة محفوظة أو إنشاء مناسبة."),
        ("إضافة مناسبة", "إدخال الاسم والتاريخ والتكرار وقاعدة اللبس والقطع والملاحظات."),
        ("التحقق من البيانات", "تعطيل الإنشاء حتى تكتمل الحقول المطلوبة وإظهار التوجيه المناسب."),
        ("مشاركة التقويم", "مشاركة خطط محددة مع شرح ما يبقى خاصًا."),
    ],
    "J": [
        ("قائمة الرحلات", "عرض الحالة الفارغة أو بطاقات الوجهات الموجودة."),
        ("الوجهة", "اختيار الوجهة والتواريخ مع دعم الإكمال التلقائي."),
        ("عدة مدن", "إضافة وجهات وتواريخ متعددة داخل الرحلة نفسها."),
        ("اختيار الأمتعة", "اختيار حقيبة يد أو حقيبة مشحونة أو الاثنين."),
        ("القطع الضرورية", "تحديد قطع مملوكة يجب إدراجها في خطة الرحلة."),
        ("الأنشطة", "اختيار أنشطة مرتبطة بالموقع أو كتابة نشاط مخصص."),
        ("جاري بناء الرحلة", "عرض تقدم التوليد ونصائح مفهومة أثناء إعداد قائمة التعبئة."),
        ("قائمة التعبئة", "تجميع القطع المملوكة وإتاحة تعديل حالة التعبئة."),
        ("القطع المقترحة", "عرض الفجوات الشرائية والمنتجات المرتبطة بسياق الرحلة."),
        ("إطلالات الرحلة", "ربط كل يوم من أيام الرحلة بإطلالة محددة."),
        ("قائمة إضافة إطلالة", "اختيار مصدر الإطلالة من قائمة التعبئة أو المحفوظات أو Muse."),
        ("اسأل Muse للرحلة", "تكوين طلب يعتمد على اليوم والوجهة والنشاط وقائمة التعبئة."),
        ("مشاركة الرحلة", "إنشاء نسخة قابلة للمشاركة من الجدول وقائمة التعبئة والإطلالات."),
        ("تعديل أو حذف الرحلة", "تعديل بيانات الرحلة أو تأكيد حذفها بإجراء واضح."),
    ],
    "K": [
        ("خلاصة Discover", "عرض شبكة كثيفة تعتمد على الصور لاكتشاف الإطلالات."),
        ("البحث الموحد", "البحث في الأشخاص والمنتجات والإطلالات والبراندات."),
        ("الفلاتر", "تصفية النتائج حسب For You أو الأكثر رواجًا أو البراندات."),
        ("تفاصيل إطلالة Discover", "عرض صاحب الإطلالة والموقع والصورة والقطع ومؤشرات التفاعل."),
        ("جرّبها عليّ", "إعادة استخدام بوابة Style Twin لعرض الإطلالة على المستخدم."),
        ("إعادة التنسيق", "نسخ إطلالة مجتمعية إلى مسودة قابلة للتعديل في الاستوديو."),
        ("المتابعة والإعجاب والتعليق", "تمثيل إجراءات المجتمع وتحديث حالتها داخل النموذج."),
        ("الإبلاغ", "اختيار سبب البلاغ وتأكيد إرساله."),
    ],
    "L": [
        ("الملف الشخصي", "عرض الصورة والاسم والمؤشرات وتاريخ العضوية ومداخل محتوى المستخدم."),
        ("استكمال الملف", "طلب استكمال اختياري يظهر في سياقه دون إيقاف الاستخدام."),
        ("الإطلالات والرحلات", "التبديل بين المحتوى المحفوظ وخطط الرحلات."),
        ("نبذة عنك", "إدارة المعلومات الشخصية التي تؤثر في التنسيق."),
        ("تفضيلات الستايل", "إدارة ما يفضله المستخدم وما لا يفضله والقواعد التي تعلمها Muse."),
        ("البراندات", "إدارة البراندات المفضلة وغير المفضلة مع البحث."),
        ("الميزانية", "تحديد نطاقات الإنفاق المناسبة لكل قطعة مقترحة."),
        ("المناسبات", "إدارة المناسبات المتكررة التي تحتاج إلى تنسيق."),
        ("Style Twin", "إدارة الصور والقياسات والوضعية والخصوصية وحد الاستخدام."),
        ("Auto-Prettify", "تشغيل أو إيقاف التحسين التلقائي للصور المرفوعة."),
        ("الإعدادات العامة", "إدارة الحساب والخصوصية والإشعارات والوحدات واللغة."),
        ("الدليل التعليمي", "شرح Muse والإضافة والتنسيق والحفظ والتغذية الراجعة."),
        ("إعدادات الإشعارات", "إدارة إشعارات التطبيق والبريد والتنظيم والتفاعل والرحلات."),
        ("مركز المساعدة", "البحث في الأدلة والتواصل مع الدعم والوصول للخصوصية والحساب."),
        ("الخصوصية والظهور", "التحكم في ظهور الملف والإطلالات وStyle Twin والبيانات وDiscover."),
    ],
    "M": [
        ("اسأل Muse", "مركز المساعد الشخصي مع اختصارات تعتمد على السياق وحقل سؤال حر."),
        ("لماذا هذه الإطلالة؟", "تفسير أسباب التوصية والانتقال إلى بدائل أو تخصيص أو حفظ الإطلالة."),
        ("فجوة الخزانة", "تحليل مبني على بيانات الخزانة يحدد النقص الحقيقي ويتجنب الشراء المكرر."),
    ],
}

SPECIAL_FLOW_EN = {
    "B-01": "Primary entry from the Closet Add button; branches to Photos (B-02), Search (B-04), or Receipt Forwarding (B-09).",
    "D-02": "Primary root screen after setup; opens Ask Muse (M-01), Style Twin (H-01), Style Studio (F-01), or Closet (C-01).",
    "F-01": "Entered from Customize/Make it mine actions, saved Looks, and Muse formalization requests.",
    "M-01": "Entered from Today, My Atelier, or the side menu; each quick prompt should preserve its intent when routing.",
    "M-02": "Entered from 'What should I wear today?', free-text Ask Muse, and 'Why this Look' actions. It currently represents a recommendation explanation.",
    "M-03": "Independent branch, not a sequential successor to M-02. Entered from 'Find a wardrobe gap' in M-01, the Wishlist insight card in C-01, or 'See the wardrobe insight' in L-01.",
}

SPECIAL_FLOW_AR = {
    "B-01": "نقطة الدخول الأساسية من زر الإضافة في Closet، وتتفرع إلى الصور B-02 أو البحث B-04 أو تمرير الإيصالات B-09.",
    "D-02": "الشاشة الرئيسية بعد الإعداد، ومنها يفتح المستخدم Ask Muse أو Style Twin أو Style Studio أو Closet.",
    "F-01": "تُفتح من إجراءات Customize أو Make it mine، ومن الإطلالات المحفوظة، ومن طلب جعل الإطلالة أكثر رسمية عبر Muse.",
    "M-01": "تُفتح من Today أو My Atelier أو السايد مينيو، ويجب أن يحافظ كل اختصار على نية المستخدم عند الانتقال.",
    "M-02": "تُفتح من سؤال ماذا أرتدي اليوم، ومن السؤال الحر، ومن إجراء Why this Look. وهي حاليًا تمثل تفسير توصية ثابتة.",
    "M-03": "فرع مستقل وليست الخطوة التالية بعد M-02. تُفتح من Find a wardrobe gap في M-01، أو كارت التحليل في Wishlist داخل C-01، أو See the wardrobe insight داخل L-01.",
}


def extract_screens():
    text = SOURCE.read_text(encoding="utf-8")
    raw_start = text.index("const raw={")
    raw_end = text.index("\n    const phaseFor", raw_start)
    raw_text = text[raw_start:raw_end]
    sections = {}
    section_matches = list(re.finditer(r"\n\s+([A-M]):\[(.*?)(?=\n\s+[A-M]:\[|\n\s+};)", raw_text, re.S))
    for match in section_matches:
        key, body = match.group(1), match.group(2)
        pairs = re.findall(r"\['([^']*)','([^']*)'\]", body)
        sections[key] = [(f"{key}-{i:02d}", title, detail) for i, (title, detail) in enumerate(pairs, 1)]
    entries = [
        ("S-00", "StyleIQ splash", "Editorial brand introduction and the first app entry point."),
        ("S-01", "Meet Muse", "Introduce the personal stylist and preview what StyleIQ can do."),
        ("S-02", "Who are we styling?", "Choose the wardrobe StyleIQ should personalize before account setup."),
    ]
    result = {"S": entries, **sections}
    assert sum(len(v) for v in result.values()) == 123, sum(len(v) for v in result.values())
    for key, values in result.items():
        assert len(AR_DATA[key]) == len(values), (key, len(AR_DATA[key]), len(values))
    return result


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name, size=None, color=None, bold=None, italic=None, rtl=False):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if rtl:
        rtl_node = r_pr.find(qn("w:rtl"))
        if rtl_node is None:
            rtl_node = OxmlElement("w:rtl")
            r_pr.append(rtl_node)


def set_paragraph_rtl(paragraph, rtl=True):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT


def add_page_field(paragraph, prefix, font_name, rtl=False):
    run = paragraph.add_run(prefix)
    set_font(run, font_name, 8, FAINT, rtl=rtl)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_styles(doc, lang):
    rtl = lang == "ar"
    font = "Arial" if rtl else "Calibri"
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    normal._element.rPr.rFonts.set(qn("w:cs"), font)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color, before, after in (
        ("Title", 28, INK, 0, 6),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 18, INK, 14, 8),
        ("Heading 2", 13, INK, 10, 5),
        ("Heading 3", 11, GOLD, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = font
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
        style._element.rPr.rFonts.set(qn("w:cs"), font)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_labeled_paragraph(doc, label, text, lang, keep=False):
    rtl = lang == "ar"
    font = "Arial" if rtl else "Calibri"
    p = doc.add_paragraph()
    set_paragraph_rtl(p, rtl)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = keep
    r1 = p.add_run(label + (" " if rtl else " "))
    set_font(r1, font, 9.5, GOLD, bold=True, rtl=rtl)
    r2 = p.add_run(text)
    set_font(r2, font, 9.5, INK, rtl=rtl)
    return p


def adjacent_text(entries, idx, lang):
    if lang == "ar":
        before = entries[idx - 1][0] if idx > 0 else "بداية القسم"
        after = entries[idx + 1][0] if idx + 1 < len(entries) else "نهاية القسم"
        return f"المجاور في دليل الشاشات: {before} قبلها و{after} بعدها. هذا ترتيب توثيقي وليس بالضرورة مسار استخدام مباشر."
    before = entries[idx - 1][0] if idx > 0 else "section start"
    after = entries[idx + 1][0] if idx + 1 < len(entries) else "section end"
    return f"Inventory adjacency: {before} before it and {after} after it. This is documentation order, not necessarily a direct user-flow dependency."


def build_document(lang, screens):
    rtl = lang == "ar"
    font = "Arial" if rtl else "Calibri"
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc, lang)

    header = section.header.paragraphs[0]
    set_paragraph_rtl(header, rtl)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    hr = header.add_run("STYLEIQ  /  SCREEN REFERENCE" if not rtl else "STYLEIQ  /  دليل الشاشات")
    set_font(hr, font, 8, MUTED, bold=True, rtl=rtl)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT if rtl else WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_rtl(footer, rtl)
    add_page_field(footer, "صفحة " if rtl else "Page ", font, rtl)

    # Editorial cover: named override on compact_reference_guide.
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    set_paragraph_rtl(kicker, rtl)
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("مرجع المنتج والتجربة" if rtl else "PRODUCT & EXPERIENCE REFERENCE")
    set_font(kr, font, 10, GOLD, bold=True, rtl=rtl)
    title = doc.add_paragraph()
    set_paragraph_rtl(title, rtl)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    tr = title.add_run("دليل شاشات StyleIQ" if rtl else "StyleIQ Screen Guide")
    set_font(tr, font, 30, INK, bold=False, rtl=rtl)
    subtitle = doc.add_paragraph(style="Subtitle")
    set_paragraph_rtl(subtitle, rtl)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("شرح وظيفة كل شاشة ومسارها داخل al.html" if rtl else "Screen-by-screen purpose and functional reference for al.html")
    set_font(sr, font, 14, MUTED, rtl=rtl)
    meta = doc.add_paragraph()
    set_paragraph_rtl(meta, rtl)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(("123 شاشة  •  إصدار 26 أغسطس 2026" if rtl else "123 screens  •  Edition: 26 August 2026"))
    set_font(mr, font, 10, FAINT, rtl=rtl)
    doc.add_paragraph()
    note = doc.add_paragraph()
    set_paragraph_rtl(note, rtl)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "يعتمد هذا الدليل على تعريفات الشاشات والتدفقات الحالية داخل al.html. ترتيب القائمة هو ترتيب Inventory، وليس دائمًا ترتيب التنقل الفعلي."
        if rtl else
        "This guide reflects the current screen definitions and flows in al.html. Inventory order is not always the same as navigation order."
    )
    set_font(nr, font, 9.5, MUTED, italic=True, rtl=rtl)
    doc.add_page_break()

    h = doc.add_paragraph("دليل الأقسام" if rtl else "Section directory", style="Heading 1")
    set_paragraph_rtl(h, rtl)
    intro = doc.add_paragraph(
        "يقسم الملف التجربة إلى 14 مجموعة وظيفية من الترحيب حتى Muse. يعرض الجدول نطاق IDs وعدد الشاشات والدور العام لكل مجموعة."
        if rtl else
        "The file groups the experience into 14 functional areas, from welcome through Muse. The table shows each ID range, screen count, and overall role."
    )
    set_paragraph_rtl(intro, rtl)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["القسم", "النطاق", "العدد", "الدور"] if rtl else ["Section", "Range", "Count", "Role"]
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, INK)
        p = cell.paragraphs[0]
        set_paragraph_rtl(p, rtl)
        r = p.add_run(value)
        set_font(r, font, 9, WHITE, bold=True, rtl=rtl)
    for key, entries in screens.items():
        cells = table.add_row().cells
        values = [SECTION_NAMES_AR[key] if rtl else SECTION_NAMES_EN[key], f"{entries[0][0]} – {entries[-1][0]}", str(len(entries)), SECTION_SHORT_AR[key] if rtl else SECTION_SHORT_EN[key]]
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            set_paragraph_rtl(p, rtl)
            r = p.add_run(value)
            set_font(r, font, 7.6, INK, rtl=rtl)
            if key in ("S", "M"):
                set_cell_shading(cells[i], GOLD_SOFT)
    set_table_geometry(table, [1550, 1350, 760, 5700])
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=45, bottom=45)

    screen_counter = 0
    for key, entries in screens.items():
        if screen_counter:
            doc.add_page_break()
        section_title = SECTION_NAMES_AR[key] if rtl else SECTION_NAMES_EN[key]
        p = doc.add_paragraph(style="Heading 1")
        set_paragraph_rtl(p, rtl)
        r = p.add_run(f"{key} — {section_title}")
        set_font(r, font, 18, INK, bold=True, rtl=rtl)
        role = doc.add_paragraph(SECTION_ROLE_AR[key] if rtl else SECTION_ROLE_EN[key])
        set_paragraph_rtl(role, rtl)
        role.paragraph_format.space_after = Pt(12)
        for idx, (screen_id, en_title, en_detail) in enumerate(entries):
            ar_title, ar_detail = AR_DATA[key][idx]
            title_text = ar_title if rtl else en_title
            detail_text = ar_detail if rtl else en_detail
            heading = doc.add_paragraph(style="Heading 2")
            set_paragraph_rtl(heading, rtl)
            heading.paragraph_format.keep_with_next = True
            badge = heading.add_run(screen_id + "  ")
            set_font(badge, font, 10, GOLD, bold=True, rtl=rtl)
            name_run = heading.add_run(title_text)
            set_font(name_run, font, 13, INK, bold=True, rtl=rtl)
            add_labeled_paragraph(doc, "الغرض:" if rtl else "Purpose:", detail_text, lang, keep=True)
            functional = (
                f"تنفذ هذه الشاشة جزءًا من وظيفة «{SECTION_NAMES_AR[key]}». {SECTION_ROLE_AR[key]}"
                if rtl else
                f"This screen implements part of the {SECTION_NAMES_EN[key]} capability. {SECTION_ROLE_EN[key]}"
            )
            add_labeled_paragraph(doc, "الوظيفة:" if rtl else "Functional role:", functional, lang, keep=True)
            add_labeled_paragraph(doc, "سياق الرحلة:" if rtl else "Journey context:", adjacent_text(entries, idx, lang), lang, keep=False)
            special = SPECIAL_FLOW_AR.get(screen_id) if rtl else SPECIAL_FLOW_EN.get(screen_id)
            if special:
                add_labeled_paragraph(doc, "ملاحظة المسار:" if rtl else "Flow note:", special, lang, keep=False)
            divider = doc.add_paragraph()
            divider.paragraph_format.space_after = Pt(2)
            p_pr = divider._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "4")
            bottom.set(qn("w:color"), LINE)
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            screen_counter += 1

    properties = doc.core_properties
    properties.title = "دليل شاشات StyleIQ" if rtl else "StyleIQ Screen Guide"
    properties.subject = "Screen-by-screen functional reference for al.html"
    properties.author = "StyleIQ Product Team"
    properties.keywords = "StyleIQ, UX, screens, product documentation, al.html"
    filename = OUT / ("StyleIQ_Screen_Guide_AR.docx" if rtl else "StyleIQ_Screen_Guide_EN.docx")
    doc.save(filename)
    return filename


if __name__ == "__main__":
    screens = extract_screens()
    outputs = [build_document("ar", screens), build_document("en", screens)]
    for path in outputs:
        print(path)
